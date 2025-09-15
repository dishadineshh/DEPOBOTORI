import warnings, logging
warnings.filterwarnings("ignore", category=UserWarning)
for noisy in ["pdfminer", "pdfplumber", "pypdfium2"]:
    logging.getLogger(noisy).setLevel(logging.ERROR)
# fetch_drive_export.py
import os, sys, re, csv, math, io, pathlib, chardet
from pathlib import Path
from dotenv import load_dotenv

# Parsers
import pdfplumber
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
from openpyxl import load_workbook

load_dotenv()

# --------- Config from .env ----------
ROOT = Path(os.getenv("DRIVE_EXPORT_DIR", "")).expanduser()
if not ROOT or not ROOT.exists():
    print(f"[drive] DRIVE_EXPORT_DIR not found: {ROOT}")
    sys.exit(1)

INCLUDE_EXT = [e.strip().lower() for e in (os.getenv("DRIVE_INCLUDE_EXT", ".pdf,.docx,.xlsx,.html")).split(",") if e.strip()]
EXCLUDE_DIRS = [d.strip().lower() for d in (os.getenv("DRIVE_EXCLUDE_DIRS", "__macosx,.git,.svn")).split(",") if d.strip()]
MAX_MB = int(os.getenv("DRIVE_MAX_FILE_MB", "40"))
DOC_CHAR_LIMIT = int(os.getenv("DRIVE_DOC_CHAR_LIMIT", "0"))  # 0 = no limit

DATA_DIR = Path(__file__).with_name("data")
OUT_CSV = DATA_DIR / "drive_export_corpus.csv"

def _skip_dir(p: Path) -> bool:
    parts = [x.lower() for x in p.parts]
    return any(ex in parts for ex in EXCLUDE_DIRS)

def _size_ok(p: Path) -> bool:
    try:
        mb = p.stat().st_size / (1024 * 1024)
        return mb <= MAX_MB
    except Exception:
        return True

def _clean_text(txt: str) -> str:
    txt = txt.replace("\r\n", "\n").replace("\r", "\n")
    txt = re.sub(r"\n{3,}", "\n\n", txt)         # collapse huge blank blocks
    txt = re.sub(r"[ \t]{2,}", " ", txt)         # collapse double spaces
    return txt.strip()

def parse_pdf(path: Path) -> str:
    with pdfplumber.open(path) as pdf:
        pages = []
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return _clean_text("\n\n".join(pages))

def parse_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    # tables too
    for table in doc.tables:
        for row in table.rows:
            text.append(" | ".join([cell.text for cell in row.cells]))
    return _clean_text("\n".join(text))

def parse_html(path: Path) -> str:
    raw = path.read_bytes()
    enc = chardet.detect(raw).get("encoding") or "utf-8"
    html = raw.decode(enc, errors="ignore")
    soup = BeautifulSoup(html, "lxml")
    # remove script/style
    for tag in soup(["script","style","noscript"]):
        tag.extract()
    txt = soup.get_text("\n")
    return _clean_text(txt)

def _try_parse_text_as_csv_bytes(path: Path) -> str:
    # Some exports are mislabeled .xlsx but are actually CSV/TSV
    raw = path.read_bytes()
    enc = chardet.detect(raw).get("encoding") or "utf-8"
    text = raw.decode(enc, errors="ignore")
    # normalize tabs to commas if it looks TSV-ish
    if "\t" in text and "," not in text:
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        rows = [ln.split("\t") for ln in lines]
        return "\n".join([" | ".join(r) for r in rows])
    return text

def parse_xlsx(path: Path) -> str:
    from openpyxl.utils.exceptions import InvalidFileException
    try:
        wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    except InvalidFileException:
        # Not a real xlsx — treat as CSV-like text
        return _clean_text(_try_parse_text_as_csv_bytes(path))
    except Exception:
        # Any other read error — best effort: try CSV-like
        return _clean_text(_try_parse_text_as_csv_bytes(path))

    lines = []
    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(h).strip() if h is not None else "" for h in rows[0]]
        for r in rows[1:]:
            parts = []
            for h, v in zip(headers, r):
                val = "" if v is None else str(v).strip()
                if val:
                    parts.append(f"{h}: {val}" if h else val)
            if parts:
                lines.append(" | ".join(parts))
    wb.close()
    return _clean_text("\n".join(lines))

def parse_file(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":   return parse_pdf(path)
    if ext == ".docx":  return parse_docx(path)
    if ext == ".html" or ext == ".htm": return parse_html(path)
    if ext == ".xlsx":  return parse_xlsx(path)
    return ""

def main():
    DATA_DIR.mkdir(exist_ok=True)
    rows = []
    file_count = 0
    picked = 0

    for p in ROOT.rglob("*"):
        if p.is_dir():
            if _skip_dir(p):
                continue
            else:
                continue
        if _skip_dir(p.parent):
            continue

        ext = p.suffix.lower()
        if ext not in INCLUDE_EXT:
            continue
        if not _size_ok(p):
            print(f"[drive] skip >{MAX_MB}MB:", p)
            continue

        file_count += 1
        try:
            txt = parse_file(p)
            if not txt:
                continue
            if DOC_CHAR_LIMIT > 0:
                txt = txt[:DOC_CHAR_LIMIT]
            src = str(p)
            rows.append({"source": src, "text": txt})
            picked += 1
            # progress
            if picked % 20 == 0:
                print(f"[drive] parsed {picked} files...")
        except Exception as e:
            print(f"[drive] error parsing {p}: {e}")

    with OUT_CSV.open("w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=["source","text"])
        w.writeheader()
        for r in rows:
            w.writerow(r)

    print(f"[drive] scanned files: {file_count}")
    print(f"[drive] parsed rows : {picked}")
    print(f"[drive] saved       : {OUT_CSV} (rows={len(rows)})")

if __name__ == "__main__":
    main()
